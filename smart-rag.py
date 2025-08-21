import os
import re
import hashlib
import tempfile
import pickle
from typing import List, Dict, Tuple, Optional, Any
from datetime import datetime
import requests
import fitz
import docx
import faiss
import nltk
import numpy as np
from rank_bm25 import BM25Okapi
import nest_asyncio
from fastapi import FastAPI, Request, HTTPException
from pyngrok import ngrok
from sentence_transformers import SentenceTransformer
import uvicorn

from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import time
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException
import logging

# Download required NLTK data for text processing
try:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
except:
    pass

class Config:
    """Configuration class containing all system settings and API credentials"""
    NGROK_AUTHTOKEN = "keep your key here"
    GROQ_API_KEY = "keep your key here"
    GROQ_URL = "https://api.groq.com/openai/v1/chat/completions"
    GROQ_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"

    # Text processing parameters
    CHUNK_SIZE = 600   # Maximum characters per text chunk
    CHUNK_OVERLAP = 100   # Overlap between consecutive chunks for context preservation
    MAX_CONTEXT_TOKENS = 6000   # Maximum context size for LLM
    DEFAULT_TOP_K = 5   # Default number of chunks to retrieve
    CONFIDENCE_THRESHOLD = 0.7   # Minimum similarity score for semantic search
    EMBEDDING_CACHE_DIR = "./embedding_cache"   # Directory to cache embeddings
    
    # Web scraping settings
    WEB_SCRAPING_TIMEOUT = 30   # Timeout for web requests
    MAX_PAGE_SIZE = 50000   # Maximum page content size to process
    USER_AGENT = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
    USE_SELENIUM_FALLBACK = True   # Use Selenium if BeautifulSoup fails

config = Config()

class EmbeddingCache:
    """Caches embeddings to avoid recomputing them for the same text"""
    
    def __init__(self):
        self.cache_dir = config.EMBEDDING_CACHE_DIR
        os.makedirs(self.cache_dir, exist_ok=True)

    def _get_cache_key(self, texts: List[str]) -> str:
        """Generate unique hash for a list of texts"""
        combined_text = "||".join(texts)
        return hashlib.md5(combined_text.encode()).hexdigest()

    def get_cached_embeddings(self, texts: List[str]) -> Optional[np.ndarray]:
        """Retrieve cached embeddings if they exist"""
        cache_key = self._get_cache_key(texts)
        cache_file = os.path.join(self.cache_dir, f"{cache_key}.npy")

        if os.path.exists(cache_file):
            try:
                return np.load(cache_file)
            except:
                return None
        return None

    def cache_embeddings(self, texts: List[str], embeddings: np.ndarray) -> None:
        """Store embeddings in cache for future use"""
        cache_key = self._get_cache_key(texts)
        cache_file = os.path.join(self.cache_dir, f"{cache_key}.npy")

        try:
            np.save(cache_file, embeddings)
        except:
            pass

class WebScraper:
    """Web scraper that uses both BeautifulSoup and Selenium for robust content extraction"""
    
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({'User-Agent': config.USER_AGENT})
        self.driver = None
        
    def _setup_selenium(self):
        """Initialize Selenium WebDriver with Chrome in headless mode"""
        if self.driver is None:
            try:
                options = Options()
                options.add_argument('--headless')   # Run without GUI
                options.add_argument('--no-sandbox')
                options.add_argument('--disable-dev-shm-usage')
                options.add_argument('--disable-gpu')
                options.add_argument('--disable-extensions')
                options.add_argument(f'--user-agent={config.USER_AGENT}')
                
                self.driver = webdriver.Chrome(options=options)
                self.driver.set_page_load_timeout(config.WEB_SCRAPING_TIMEOUT)
                
            except Exception as e:
                print(f"Warning: Selenium setup failed: {e}")
                self.driver = None
        
        return self.driver is not None

    def _extract_text_with_beautifulsoup(self, html_content: str, url: str) -> str:
        """Extract clean text from HTML using BeautifulSoup"""
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove unwanted elements that don't contain useful content
            for element in soup(['script', 'style', 'nav', 'header', 'footer', 
                               'sidebar', 'advertisement', 'ads', 'menu']):
                element.decompose()
            
            # Try to find main content areas first
            content_selectors = [
                'main', 'article', '[role="main"]', '.content', '.main-content',
                '.post-content', '.entry-content', '.article-content', '.page-content'
            ]
            
            main_content = None
            for selector in content_selectors:
                main_content = soup.select_one(selector)
                if main_content:
                    break
            
            # Fallback to entire body if no main content found
            if not main_content:
                main_content = soup.find('body') or soup
            
            text_parts = []
            
            # Extract title
            title = soup.find('title')
            if title:
                text_parts.append(f"TITLE: {title.get_text().strip()}")
            
            # Extract text from relevant HTML elements
            for element in main_content.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'section']):
                text = element.get_text().strip()
                if text and len(text) > 10:  # Only include meaningful text
                    text_parts.append(text)
            
            full_text = '\n\n'.join(text_parts)
            
            # Clean up excessive whitespace
            full_text = re.sub(r'\n\s*\n\s*\n', '\n\n', full_text)
            full_text = re.sub(r'\s+', ' ', full_text)
            
            return full_text.strip()
            
        except Exception as e:
            print(f"BeautifulSoup extraction error: {e}")
            return ""

    def _scrape_with_requests(self, url: str) -> str:
        """Primary scraping method using requests and BeautifulSoup"""
        try:
            response = self.session.get(url, timeout=config.WEB_SCRAPING_TIMEOUT)
            response.raise_for_status()
            
            # Check if content is HTML
            content_type = response.headers.get('content-type', '').lower()
            if 'html' not in content_type:
                if 'text' in content_type:
                    return response.text[:config.MAX_PAGE_SIZE]
                else:
                    raise ValueError(f"Unsupported content type: {content_type}")
            
            return self._extract_text_with_beautifulsoup(response.text, url)
            
        except Exception as e:
            print(f"Requests scraping failed for {url}: {e}")
            return ""

    def _scrape_with_selenium(self, url: str) -> str:
        """Fallback scraping method using Selenium for JavaScript-heavy sites"""
        if not self._setup_selenium():
            return ""
        
        try:
            self.driver.get(url)
            
            # Wait for page to load
            WebDriverWait(self.driver, 10).until(
                EC.presence_of_element_located((By.TAG_NAME, "body"))
            )
            
            # Additional wait for dynamic content
            time.sleep(3)
            
            html_content = self.driver.page_source
            return self._extract_text_with_beautifulsoup(html_content, url)
            
        except Exception as e:
            print(f"Selenium scraping failed for {url}: {e}")
            return ""

    def scrape_website(self, url: str) -> str:
        """Main scraping method with fallback strategy"""
        print(f"Scraping website: {url}")
        
        # Try requests first (faster)
        content = self._scrape_with_requests(url)
        
        # Fallback to Selenium if needed
        if not content and config.USE_SELENIUM_FALLBACK:
            print("Trying Selenium fallback...")
            content = self._scrape_with_selenium(url)
        
        if not content:
            raise ValueError(f"Unable to extract content from {url}")
        
        # Truncate if too large
        if len(content) > config.MAX_PAGE_SIZE:
            content = content[:config.MAX_PAGE_SIZE] + "\n\n[Content truncated due to size limits]"
        
        print(f"Extracted {len(content)} characters from website")
        return content

    def __del__(self):
        """Cleanup Selenium driver when object is destroyed"""
        if self.driver:
            try:
                self.driver.quit()
            except:
                pass

class EnhancedDocumentLoader:
    """Handles loading and processing of various document types (PDFs, DOCX, websites)"""
    
    def __init__(self, text_processor):
        self.text_processor = text_processor
        self.file_cache = {}  # Cache downloaded files
        self.web_scraper = WebScraper()

    def _is_web_url(self, url: str) -> bool:
        """Determine if URL is a website or a direct file link"""
        parsed = urlparse(url.lower())
        
        # Check if URL path ends with file extension
        path = parsed.path.lower()
        file_extensions = ['.pdf', '.docx', '.doc', '.txt', '.rtf']
        
        if any(path.endswith(ext) for ext in file_extensions):
            return False
        
        # If it's HTTP/HTTPS and not a file, treat as website
        if parsed.scheme in ['http', 'https']:
            return True
        
        return False

    def download_file(self, url: str) -> str:
        """Download file from URL and cache it locally"""
        url_hash = hashlib.md5(url.encode()).hexdigest()

        # Return cached file if exists
        if url_hash in self.file_cache:
            return self.file_cache[url_hash]

        headers = {'User-Agent': config.USER_AGENT}
        resp = requests.get(url, headers=headers, timeout=30)
        resp.raise_for_status()

        # Determine file extension
        file_ext = self._get_file_extension(url, resp.headers.get('content-type', ''))
        temp_path = tempfile.mktemp(suffix=file_ext)

        # Save file to temporary location
        with open(temp_path, "wb") as f:
            f.write(resp.content)

        self.file_cache[url_hash] = temp_path
        return temp_path

    def _get_file_extension(self, url: str, content_type: str = '') -> str:
        """Determine file extension from URL or content type"""
        url_clean = url.split('?')[0].lower()

        # Check URL extension first
        if url_clean.endswith('.pdf'):
            return '.pdf'
        elif url_clean.endswith('.docx') or url_clean.endswith('.doc'):
            return '.docx'

        # Check content type
        if content_type:
            if 'pdf' in content_type.lower():
                return '.pdf'
            elif 'word' in content_type.lower() or 'docx' in content_type.lower():
                return '.docx'

        # Default to PDF
        return '.pdf'

    def extract_text_from_pdf(self, path: str) -> str:
        """Extract text from PDF using PyMuPDF"""
        text_content = ""
        with fitz.open(path) as doc:
            for page in doc:
                try:
                    page_text = page.get_text("text")
                    if page_text.strip():
                        text_content += f"\n{page_text}"
                except:
                    continue

        if not text_content.strip():
            raise ValueError("No text content could be extracted from the PDF")

        return self.text_processor.clean_text(text_content)

    def extract_text_from_docx(self, path: str) -> str:
        """Extract text from DOCX using python-docx"""
        doc = docx.Document(path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts)
        return self.text_processor.clean_text(full_text)

    def load_document(self, url: str) -> str:
        """Main document loading method that handles different content types"""
        if self._is_web_url(url):
            print("Detected website URL, scraping content...")
            web_content = self.web_scraper.scrape_website(url)
            return self.text_processor.clean_text(web_content)
        
        else:
            print("Detected file URL, downloading and extracting...")
            path = self.download_file(url)
            file_extension = path.lower().split('.')[-1] if '.' in path else 'pdf'

            # Try appropriate extractor based on file type
            if file_extension == 'pdf' or 'pdf' in url.lower():
                return self.extract_text_from_pdf(path)
            elif file_extension in ['docx', 'doc'] or any(x in url.lower() for x in ['docx', 'doc']):
                return self.extract_text_from_docx(path)
            else:
                # Try PDF first, then DOCX as fallback
                try:
                    return self.extract_text_from_pdf(path)
                except:
                    return self.extract_text_from_docx(path)

class TextProcessor:
    """Handles text cleaning and chunking operations"""
    
    def __init__(self):
        # Initialize stopwords for text processing
        try:
            from nltk.corpus import stopwords
            self.stop_words = set(stopwords.words('english'))
        except:
            # Fallback stopwords if NLTK download fails
            self.stop_words = set(['the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'])

    def clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        if not text or not isinstance(text, str):
            return ""
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text.strip())
        # Remove unwanted characters but keep punctuation
        text = re.sub(r'[^\w\s\.\!\?\-\,\:\;]', '', text)
        return text

    def fast_chunking(self, text: str) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks for better retrieval"""
        if not text:
            return []

        # Split into sentences for better chunk boundaries
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        chunk_id = 0
        current_chunk = ""

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            # Check if adding sentence exceeds chunk size
            if len(current_chunk + sentence) > config.CHUNK_SIZE:
                # Save current chunk if it's meaningful
                if current_chunk and len(current_chunk.strip()) > 50:
                    chunks.append({
                        'id': chunk_id,
                        'text': current_chunk.strip()
                    })
                    chunk_id += 1

                # Create overlap with previous chunk
                if config.CHUNK_OVERLAP > 0 and current_chunk:
                    overlap_text = current_chunk[-config.CHUNK_OVERLAP:] if len(current_chunk) > config.CHUNK_OVERLAP else current_chunk
                    current_chunk = overlap_text + " " + sentence
                else:
                    current_chunk = sentence
            else:
                current_chunk += " " + sentence if current_chunk else sentence

        # Add final chunk
        if current_chunk and len(current_chunk.strip()) > 50:
            chunks.append({
                'id': chunk_id,
                'text': current_chunk.strip()
            })

        return chunks

class SmartRetriever:
    """Hybrid retrieval system using both semantic similarity and keyword matching"""
    
    def __init__(self, embedder: SentenceTransformer, embedding_cache: EmbeddingCache):
        self.embedder = embedder
        self.embedding_cache = embedding_cache
        self.faiss_index = None   # For semantic search
        self.bm25_index = None    # For keyword search
        self.chunks = []

    def build_indices(self, chunks: List[Dict[str, Any]]) -> None:
        """Build both semantic and keyword search indices"""
        self.chunks = chunks
        chunk_texts = [chunk['text'] for chunk in chunks]

        # Try to get cached embeddings first
        embeddings = self.embedding_cache.get_cached_embeddings(chunk_texts)

        if embeddings is None:
            # Generate embeddings if not cached
            embeddings = self.embedder.encode(
                chunk_texts,
                convert_to_tensor=False,
                show_progress_bar=False,
                batch_size=32
            )
            # Cache for future use
            self.embedding_cache.cache_embeddings(chunk_texts, embeddings)

        # Build FAISS index for semantic similarity search
        dimension = embeddings.shape[1]
        self.faiss_index = faiss.IndexFlatIP(dimension)   # Inner product for cosine similarity
        faiss.normalize_L2(embeddings)   # Normalize for cosine similarity
        self.faiss_index.add(embeddings)

        # Build BM25 index for keyword search
        tokenized_chunks = [chunk['text'].lower().split() for chunk in chunks]
        self.bm25_index = BM25Okapi(tokenized_chunks)

    def smart_retrieve_context(self, query: str, top_k: int = None) -> str:
        """Retrieve relevant context using hybrid approach"""
        top_k = top_k or config.DEFAULT_TOP_K
        
        # Get semantic search results
        semantic_results = self._semantic_search(query, top_k)

        # If semantic search has high confidence, use it primarily
        if semantic_results:
            max_semantic_score = max(score for _, score in semantic_results)
            if max_semantic_score > config.CONFIDENCE_THRESHOLD:
                return self._assemble_context(semantic_results)

        # Otherwise, combine with keyword search
        keyword_results = self._keyword_search(query, top_k)

        if keyword_results:
            max_keyword_score = max(score for _, score in keyword_results)
            if max_keyword_score > 5.0:  # BM25 threshold
                return self._combine_results(semantic_results, keyword_results, top_k)

        # Fallback to semantic results
        return self._assemble_context(semantic_results)

    def _semantic_search(self, query: str, k: int) -> List[Tuple[int, float]]:
        """Perform semantic similarity search using embeddings"""
        query_embedding = self.embedder.encode([query], convert_to_tensor=False)
        faiss.normalize_L2(query_embedding)
        similarities, indices = self.faiss_index.search(query_embedding, min(k, len(self.chunks)))
        return [(int(idx), float(sim)) for idx, sim in zip(indices[0], similarities[0]) if idx != -1]

    def _keyword_search(self, query: str, k: int) -> List[Tuple[int, float]]:
        """Perform keyword-based search using BM25"""
        query_tokens = query.lower().split()
        scores = self.bm25_index.get_scores(query_tokens)
        top_indices = np.argsort(scores)[::-1][:min(k, len(scores))]
        return [(int(idx), float(scores[idx])) for idx in top_indices if scores[idx] > 0]

    def _combine_results(self, semantic_results: List[Tuple[int, float]],
                        keyword_results: List[Tuple[int, float]], top_k: int) -> str:
        """Combine semantic and keyword search results with weighted scoring"""
        semantic_weight = 0.7
        keyword_weight = 0.3

        # Convert to dictionaries for easier manipulation
        sem_scores = {idx: score for idx, score in semantic_results}
        key_scores = {idx: score for idx, score in keyword_results}

        # Normalize semantic scores
        if sem_scores:
            max_sem = max(sem_scores.values())
            min_sem = min(sem_scores.values())
            if max_sem > min_sem:
                sem_scores = {idx: (score - min_sem) / (max_sem - min_sem) for idx, score in sem_scores.items()}

        # Normalize keyword scores
        if key_scores:
            max_key = max(key_scores.values())
            min_key = min(key_scores.values())
            if max_key > min_key:
                key_scores = {idx: (score - min_key) / (max_key - min_key) for idx, score in key_scores.items()}

        # Combine scores
        all_indices = set(sem_scores.keys()) | set(key_scores.keys())
        combined_scores = []

        for idx in all_indices:
            sem_score = sem_scores.get(idx, 0)
            key_score = key_scores.get(idx, 0)
            combined_score = semantic_weight * sem_score + keyword_weight * key_score
            combined_scores.append((idx, combined_score))

        # Sort by combined score and return top results
        combined_scores.sort(key=lambda x: x[1], reverse=True)
        top_results = combined_scores[:top_k]
        return self._assemble_context(top_results)

    def _assemble_context(self, results: List[Tuple[int, float]]) -> str:
        """Assemble context from retrieved chunks, respecting token limits"""
        if not results:
            return ''

        context_parts = []
        total_chars = 0

        for idx, score in results:
            chunk_text = self.chunks[idx]['text']
            # Stop if adding this chunk would exceed context limit
            if total_chars + len(chunk_text) > config.MAX_CONTEXT_TOKENS:
                break
            context_parts.append(chunk_text)
            total_chars += len(chunk_text)

        return "\n\n".join(context_parts)

class AnswerGenerator:
    """Generates answers using the Groq API with optimized batch processing"""
    
    def __init__(self):
        # Setup HTTP session with connection pooling
        self.session = requests.Session()
        adapter = requests.adapters.HTTPAdapter(pool_connections=10, pool_maxsize=20, max_retries=3)
        self.session.mount('https://', adapter)
        self.session.mount('http://', adapter)

        self.session.headers.update({
            "Authorization": f"Bearer {config.GROQ_API_KEY}",
            "Content-Type": "application/json"
        })

    def generate_answers(self, questions_with_context: List[Dict[str, Any]]) -> List[str]:
        """Generate answers for multiple questions in a single API call"""
        return self._generate_single_batch_call(questions_with_context)

    def _generate_single_batch_call(self, questions_with_context: List[Dict[str, Any]]) -> List[str]:
        """Batch all questions into a single LLM call for efficiency"""
        # System prompt with strict formatting instructions
        system_prompt = """You are an expert assistant. Please search for the answers thoroughly. Go through the documents attentively and carefully for the answers. For each question, provide a clear, concise answer (1-2 sentences max) based strictly on the provided context. Convert technical terms to user-friendly language. Make answers interpretable.

CRITICAL FORMATTING RULES:
1. Format your response EXACTLY as numbered answers with NO introductory text. But every answer must be at least one sentence, not single word or two words.
2. Use the format: 1) [Answer to question 1]
3. Each answer should be on its own line
4. Do NOT include phrases like "Here are the answers:" or any other introductory text
5. Start directly with "1) [your answer]"
6. If it is a yes or no question, your answer must start with either yes or no

Example format:
1) The answer to the first question based on the context.
2) The answer to the second question based on the context.
3) The answer to the third question based on the context."""

        # Build user prompt with all questions and contexts
        prompt_parts = []
        for idx, qc in enumerate(questions_with_context, 1):
            # Limit context size to prevent token overflow
            limited_context = qc['context'][:2000]
            prompt_parts.append(f"""
Question {idx}: {qc['question']}
Context {idx}: {limited_context}
""")

        user_prompt = f"""Answer each question based on its corresponding context:
{''.join(prompt_parts)}

Provide ONLY numbered answers with NO introductory text, starting directly with "1)":"""

        # Call API and parse response
        response_text = self._call_groq_api(system_prompt, user_prompt)
        return self._parse_batch_response(response_text, len(questions_with_context))

    def _call_groq_api(self, system_prompt: str, user_prompt: str) -> str:
        """Make API call to Groq with error handling"""
        payload = {
            "model": config.GROQ_MODEL,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": 0.1,  # Low temperature for consistent formatting
            "max_tokens": 1200,
            "top_p": 0.9
        }

        response = self.session.post(config.GROQ_URL, json=payload, timeout=60)
        response.raise_for_status()

        response_data = response.json()
        content = response_data["choices"][0]["message"]["content"].strip()

        # Remove common unwanted prefixes
        unwanted_prefixes = [
            "here are the answers:",
            "the answers are:",
            "below are the answers:",
            "here are your answers:",
            "answers:",
        ]

        content_lower = content.lower()
        for prefix in unwanted_prefixes:
            if content_lower.startswith(prefix):
                content = content[len(prefix):].strip()
                break

        return content

    def _parse_batch_response(self, batch_response: str, expected_count: int) -> List[str]:
        """Parse the batch response into individual answers"""
        try:
            import re
            lines = batch_response.strip().split('\n')
            clean_lines = []

            # Remove unwanted lines
            for line in lines:
                line = line.strip()
                if not line or line.lower().startswith(('here are', 'the answers', 'below are', 'answers:')):
                    continue
                clean_lines.append(line)

            cleaned_response = '\n'.join(clean_lines)
            
            # Split by numbered patterns
            pattern = r'^\d+\)\s*'
            parts = re.split(pattern, cleaned_response, flags=re.MULTILINE)

            # Remove empty first part if exists
            if parts and not parts[0].strip():
                parts = parts[1:]

            answers = []
            for i, part in enumerate(parts[:expected_count]):
                # Clean each answer
                cleaned_answer = part.strip()
                cleaned_answer = re.sub(r'^\d+[\)\.]?\s*', '', cleaned_answer).strip()

                if cleaned_answer:
                    # Take first line for cleaner answers
                    first_part = cleaned_answer.split('\n')[0].strip()
                    answers.append(first_part if first_part else cleaned_answer)
                else:
                    answers.append(f"Unable to generate answer for question {i+1}")

            # Use fallback parsing if primary method fails
            if len(answers) < expected_count:
                answers = self._fallback_parse_method(batch_response, expected_count)

            # Ensure we have the right number of answers
            while len(answers) < expected_count:
                answers.append("Answer not available")

            return answers[:expected_count]

        except Exception as e:
            print(f"Error parsing batch response: {e}")
            return self._fallback_parse_method(batch_response, expected_count)

    def _fallback_parse_method(self, batch_response: str, expected_count: int) -> List[str]:
        """Fallback method for parsing responses when primary method fails"""
        try:
            import re
            lines = batch_response.strip().split('\n')
            answers = []

            # Look for numbered lines first
            for line in lines:
                line = line.strip()
                if re.match(r'^\d+\)', line):
                    answer = re.sub(r'^\d+\)\s*', '', line).strip()
                    if answer and not answer.lower().startswith(('here are', 'the answers', 'below are')):
                        answers.append(answer)

            # Try regex extraction if not enough answers
            if len(answers) < expected_count:
                all_text = ' '.join(lines)
                numbered_parts = re.findall(r'\d+\)\s*([^0-9]+?)(?=\d+\)|$)', all_text)

                for part in numbered_parts:
                    cleaned = part.strip()
                    if cleaned and len(answers) < expected_count:
                        answers.append(cleaned)

            # Final fallback: split by sentences
            if len(answers) < expected_count:
                sentences = re.split(r'[.!?]+', batch_response)
                relevant_sentences = [s.strip() for s in sentences if s.strip() and len(s.strip()) > 10]

                for sentence in relevant_sentences:
                    if len(answers) < expected_count:
                        answers.append(sentence)

            # Pad with error messages if still not enough
            while len(answers) < expected_count:
                answers.append(f"Answer not available for question {len(answers) + 1}")

            return answers[:expected_count]

        except:
            # Ultimate fallback
            return [f"Error processing answer {i+1}" for i in range(expected_count)]

class EnhancedRAGSystem:
    """Main RAG system that orchestrates all components"""
    
    def __init__(self):
        import torch
        # Use GPU if available for faster embedding generation
        device = 'cuda' if torch.cuda.is_available() else 'cpu'

        # Initialize all components
        self.embedder = SentenceTransformer("multi-qa-MiniLM-L6-cos-v1", device=device)
        self.embedding_cache = EmbeddingCache()
        self.text_processor = TextProcessor()
        self.document_loader = EnhancedDocumentLoader(self.text_processor)
        self.retriever = SmartRetriever(self.embedder, self.embedding_cache)
        self.answer_generator = AnswerGenerator()

        # System state
        self.document_cache = {}  # Cache processed documents
        self.is_initialized = False

    def process_document(self, doc_url: str) -> None:
        """Process a document (PDF, DOCX, or website) and prepare it for querying"""
        # Check if document is already processed
        doc_hash = hashlib.md5(doc_url.encode()).hexdigest()

        if doc_hash in self.document_cache:
            # Use cached document
            cached_data = self.document_cache[doc_hash]
            self.retriever.build_indices(cached_data['chunks'])
            self.is_initialized = True
            return

        # Load and process new document
        doc_text = self.document_loader.load_document(doc_url)
        chunks = self.text_processor.fast_chunking(doc_text)

        if not chunks:
            raise HTTPException(status_code=400, detail="No valid content found in document")

        # Build search indices
        self.retriever.build_indices(chunks)
        
        # Cache for future use
        self.document_cache[doc_hash] = {'chunks': chunks}
        self.is_initialized = True

    def answer_questions(self, questions: List[str]) -> List[str]:
        """Answer multiple questions using the processed document"""
        if not self.is_initialized:
            raise HTTPException(status_code=400, detail="No document has been processed yet")

        # Retrieve relevant context for each question
        questions_with_context = []
        for question in questions:
            context = self.retriever.smart_retrieve_context(question)
            questions_with_context.append({
                'question': question,
                'context': context
            })

        # Generate answers in a single batch call
        return self.answer_generator.generate_answers(questions_with_context)

# Initialize FastAPI app and RAG system
app = FastAPI()
rag_system = EnhancedRAGSystem()

# Traditional Q&A endpoint
@app.post("/v1/api")
async def v1_api(request: Request):
    """Handle traditional question-answering requests"""
    try:
        data = await request.json()

        # Validate input
        if not data.get("documents") or not data.get("questions"):
            raise HTTPException(status_code=400, detail="Both 'documents' and 'questions' are required")

        doc_url = data["documents"]
        questions = data["questions"]

        if not isinstance(questions, list) or not questions:
            raise HTTPException(status_code=400, detail="'questions' must be a non-empty list")

        # Process document and answer questions
        print(f"Processing: {doc_url}")
        rag_system.process_document(doc_url)

        print(f"Answering {len(questions)} questions...")
        answers = rag_system.answer_questions(questions)

        return {"answers": answers}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

# Root endpoint with API documentation
@app.get("/")
async def root():
    """API documentation and feature overview"""
    return {
        "message": "Enhanced RAG API with Web Scraping",
        "features": [
            "Websites (HTML scraping with BeautifulSoup + Selenium)",
            "PDF documents",
            "DOCX documents",
            "Single LLM call optimization",
            "Smart hybrid retrieval (semantic + keyword)",
            "Embedding caching"
        ],
        "endpoints": {
            "/v1/api": {
                "method": "POST",
                "description": "Traditional Q&A with documents",
                "example_payload": {
                    "documents": "https://example.com or file.pdf",
                    "questions": ["What is the main topic?"]
                }
            }
        }
    }

def install_dependencies():
    """Install required dependencies that might be missing"""
    import subprocess
    import sys
    
    packages = [
        "beautifulsoup4",
        "selenium",
        "webdriver-manager"
    ]
    
    for package in packages:
        try:
            __import__(package.replace("-", "_"))
            print(f"{package} already installed")
        except ImportError:
            print(f"Installing {package}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])

def start_server():
    """Start the FastAPI server with ngrok tunnel"""
    try:
        # Install any missing dependencies
        install_dependencies()
        
        # Setup ngrok tunnel
        ngrok.set_auth_token(config.NGROK_AUTHTOKEN)
        nest_asyncio.apply()

        tunnel = ngrok.connect(8000)
        public_url = tunnel.public_url

        # Print server information
        print("ENHANCED RAG API LIVE:", public_url)
        print("FEATURES:")
        print("    Website scraping (BeautifulSoup + Selenium)")
        print("    PDF & DOCX documents")
        print("    Single LLM call optimization")
        print("ENDPOINTS:")
        print(f"    Q&A: POST {public_url}/v1/api")
        print("Example Usage:")
        print("    Websites: https://en.wikipedia.org/wiki/AI")
        print("    Documents: https://example.com/document.pdf")
        print(f"Model: {config.GROQ_MODEL}")
        print("Embedder: multi-qa-MiniLM-L6-cos-v1")

        # Start the server
        uvicorn.run(app, host="0.0.0.0", port=8000, log_level="error")
        
    except Exception as e:
        print(f"Error starting server: {e}")
        print("Make sure you have Chrome/Chromium installed for Selenium")

# Main execution
if __name__ == "__main__":
    start_server()
